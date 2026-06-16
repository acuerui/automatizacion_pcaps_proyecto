from __future__ import annotations

import sys
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def text_of(node: Tag) -> str:
    return " ".join(node.get_text(" ", strip=True).split())


def add_paragraph_with_inline(doc: Document, node: Tag, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    append_inline(paragraph, node)
    return paragraph


def append_inline(paragraph, node):
    for child in node.children:
        if isinstance(child, NavigableString):
            paragraph.add_run(str(child))
        elif isinstance(child, Tag):
            if child.name == "br":
                paragraph.add_run().add_break()
            elif child.name in {"strong", "b"}:
                run = paragraph.add_run(text_of(child))
                run.bold = True
            elif child.name in {"em", "i"}:
                run = paragraph.add_run(text_of(child))
                run.italic = True
            elif child.name == "code":
                run = paragraph.add_run(child.get_text(strip=True))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            elif child.name == "a":
                run = paragraph.add_run(text_of(child))
                run.underline = True
            else:
                append_inline(paragraph, child)


def add_table(doc: Document, table_tag: Tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    max_cols = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"], recursive=False)
        for c_idx, cell in enumerate(cells):
            out_cell = table.cell(r_idx, c_idx)
            out_cell.text = ""
            paragraph = out_cell.paragraphs[0]
            append_inline(paragraph, cell)
            if cell.name == "th":
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph()


def add_list(doc: Document, list_tag: Tag, ordered: bool = False):
    style = "List Number" if ordered else "List Bullet"
    for item in list_tag.find_all("li", recursive=False):
        add_paragraph_with_inline(doc, item, style=style)


def iter_content(root: Tag):
    for child in root.children:
        if isinstance(child, NavigableString):
            continue
        if not isinstance(child, Tag):
            continue
        if child.name in {"h1", "h2", "h3", "p", "div", "table", "ul", "ol"}:
            yield child


def convert_html_to_docx(html_path: Path, docx_path: Path):
    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")
    root = soup.find("main") or soup.body or soup

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    for node in iter_content(root):
        if node.name == "h1":
            p = doc.add_heading(text_of(node), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif node.name == "h2":
            doc.add_heading(text_of(node), level=1)
        elif node.name == "h3":
            doc.add_heading(text_of(node), level=2)
        elif node.name == "p":
            if text_of(node):
                add_paragraph_with_inline(doc, node)
        elif node.name == "div":
            for nested in iter_content(node):
                if nested.name == "p" and text_of(nested):
                    add_paragraph_with_inline(doc, nested)
                elif nested.name in {"strong", "b"}:
                    p = doc.add_paragraph()
                    append_inline(p, nested)
            direct_text = text_of(node)
            if direct_text and not node.find(["p", "table", "ul", "ol", "h1", "h2", "h3"]):
                add_paragraph_with_inline(doc, node)
        elif node.name == "table":
            add_table(doc, node)
        elif node.name == "ul":
            add_list(doc, node, ordered=False)
        elif node.name == "ol":
            add_list(doc, node, ordered=True)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: python export_html_to_docx.py <html_dir> <docx_dir>")
        return 2
    html_dir = Path(sys.argv[1])
    docx_dir = Path(sys.argv[2])
    html_files = sorted(p for p in html_dir.glob("*.html") if p.is_file())
    if not html_files:
        print(f"No HTML files found in {html_dir}")
        return 1
    for html_file in html_files:
        out = docx_dir / f"{html_file.stem}.docx"
        convert_html_to_docx(html_file, out)
        print(f"{html_file.name} -> {out.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
